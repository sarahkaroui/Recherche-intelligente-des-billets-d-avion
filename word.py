from docx import Document
import re

def clean_llm_text(text: str) -> str:
    """Supprime les caractères parasites et espaces inutiles."""
    if not text:
        return ""
    text = text.replace("\xa0", " ")  
    text = re.sub(r"\s+", " ", text)  
    return text.strip()

def save_llm_to_word(answer: str, path: str = "flight_agent_result.docx"):
    """Écrase le fichier Word et ajoute la réponse LLM nettoyée."""
    doc = Document()  
    doc.add_paragraph("Résultat de l'agent de vols").bold = True
    doc.add_paragraph(clean_llm_text(answer))
    doc.save(path)
    return path
